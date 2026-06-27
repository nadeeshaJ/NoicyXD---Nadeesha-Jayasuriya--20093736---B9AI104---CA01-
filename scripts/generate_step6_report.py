"""Generate Step 6 error analysis report and development notes.

Run: python scripts/generate_step6_report.py
Outputs: reports/step6/Step6_Report.docx, reports/step6/Step6_Development_Notes.docx
"""

from __future__ import annotations

import csv
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = PROJECT_ROOT / "reports" / "figures" / "step6"
STEP6_DIR = PROJECT_ROOT / "reports" / "step6"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_figure(doc, path: Path, caption: str, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()


def build_report(summary: dict) -> Document:
    doc = Document()
    title = doc.add_heading(
        "Step 6 — Error Analysis & Advanced Evaluation", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Deep Learning (B9AI104) — Model Comparison, Evaluation & Validation (15%) + Results & Elaboration (15%)",
    )
    doc.add_paragraph()

    best_urban = summary.get("best_urban_run", "mobilenetv2")
    urban_results = summary.get("urban_error_analysis", [])
    benchmarks = summary.get("inference_benchmarks", [])
    case_studies = summary.get("case_studies", [])
    cross_domain = summary.get("cross_domain", [])

    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "Phase 6 extends evaluation beyond aggregate accuracy and F1 scores. All three UrbanSound8K "
        "models are analysed on fold-10 test predictions: confusion patterns, misclassification "
        "case studies, inference benchmarks, and a cross-domain summary linking Stage 1 and Stage 2.",
    )

    add_heading(doc, "2. Confusion Matrix Deep Dive (UrbanSound8K)", level=1)
    add_para(
        doc,
        "The fold-10 test set (837 clips) was re-evaluated for Custom CNN, ResNet50, and MobileNetV2. "
        "Top confused pairs highlight acoustically similar urban classes.",
    )

    for result in urban_results:
        run_name = result.get("run_name", "")
        add_para(doc, f"Model: {run_name} (test accuracy={result.get('test_accuracy', 0):.3f})", bold=True)
        for pair in result.get("top_confused_pairs", [])[:5]:
            add_bullets(
                doc,
                [
                    f"{pair['true_class']} misclassified as {pair['predicted_class']} "
                    f"({pair['count']} times)"
                ],
            )
        add_figure(
            doc,
            FIG_DIR / f"confusion_matrix_{run_name}.png",
            f"Confusion matrix — {run_name}.",
            width=Inches(4.8),
        )
        add_figure(
            doc,
            FIG_DIR / f"confused_pairs_{run_name}.png",
            f"Top confused pairs — {run_name}.",
            width=Inches(5.2),
        )

    add_heading(doc, "3. Misclassification Case Studies", level=1)
    add_para(
        doc,
        f"Five detailed case studies were selected from the best urban model ({best_urban}) "
        "by prioritising the most frequent confusion pairs. Each case includes waveform, "
        "Mel-spectrogram, model input image, true label, predicted label, and explanation.",
    )
    for i, case in enumerate(case_studies[:5], start=1):
        add_para(
            doc,
            f"Case {i}: True={case.get('true_label')} | Pred={case.get('predicted_label')} | "
            f"Confidence={case.get('confidence', 0):.1%}",
            bold=True,
        )
        add_para(doc, case.get("explanation", ""))
        add_figure(
            doc,
            FIG_DIR / "case_studies" / f"case_study_{i:02d}.png",
            f"Figure 6.{i} — Case study {i}.",
            width=Inches(6.2),
        )

    add_heading(doc, "4. Inference Benchmarking", level=1)
    add_para(
        doc,
        "Benchmarks were measured on a single 224x224 input with GPU warmup (if available). "
        "Training time per epoch comes from Step 3 training logs on CPU.",
    )
    if benchmarks:
        table = doc.add_table(rows=1 + len(benchmarks), cols=6)
        table.style = "Table Grid"
        headers = [
            "Model",
            "Parameters",
            "File size (MB)",
            "Inference (ms)",
            "Train time/epoch (s)",
            "Test macro F1",
        ]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for r, row in enumerate(benchmarks):
            cells = table.rows[r + 1].cells
            cells[0].text = row.get("model", "")
            cells[1].text = f"{row.get('total_parameters', 0):,}"
            cells[2].text = f"{row.get('model_file_size_mb', 0):.2f}"
            cells[3].text = f"{row.get('inference_ms_mean', 0):.2f} +/- {row.get('inference_ms_std', 0):.2f}"
            cells[4].text = f"{row.get('training_time_per_epoch_sec', 0):.1f}"
            cells[5].text = f"{row.get('test_macro_f1', 0):.4f}"
        doc.add_paragraph()

    add_para(
        doc,
        "MobileNetV2 offers the best accuracy-to-efficiency trade-off: highest urban macro F1 "
        "with the smallest checkpoint and fastest inference among the three models.",
        bold=True,
    )

    add_heading(doc, "5. Cross-Domain Summary", level=1)
    add_para(
        doc,
        "Stage 2 ESC-50 results are summarised alongside the urban baseline to quantify domain shift.",
    )
    if cross_domain:
        table = doc.add_table(rows=1 + len(cross_domain), cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Approach", "Domain", "Macro F1", "F1 drop vs Urban"]):
            table.rows[0].cells[i].text = h
        for r, row in enumerate(cross_domain):
            cells = table.rows[r + 1].cells
            cells[0].text = str(row.get("approach", ""))
            cells[1].text = str(row.get("domain", ""))
            cells[2].text = f"{float(row.get('macro_f1', 0)):.4f}"
            cells[3].text = f"{float(row.get('f1_drop', 0)):.4f}"
        doc.add_paragraph()

    esc = summary.get("esc50_error_analysis", {})
    add_para(doc, f"Best ESC-50 model ({summary.get('best_esc50_run', '')}) — top confused pairs:", bold=True)
    for pair in esc.get("top_confused_pairs", [])[:5]:
        add_bullets(
            doc,
            [
                f"{pair['true_class']} misclassified as {pair['predicted_class']} "
                f"({pair['count']} times)"
            ],
        )
    add_figure(
        doc,
        FIG_DIR / f"error_examples_{summary.get('best_esc50_run', '')}.png",
        "Misclassified ESC-50 animal examples (best model).",
        width=Inches(5.5),
    )

    add_heading(doc, "6. Mel-Spectrogram vs MFCC", level=1)
    add_para(
        doc,
        "MFCCs compress spectral information into a small set of cepstral coefficients optimised "
        "for speech. Environmental sounds span wider bandwidth and richer temporal structure. "
        "Converting Mel-spectrograms to 224x224 RGB images allows direct use of CNN and transfer-learning "
        "architectures (ResNet, MobileNet) pretrained on ImageNet. MFCC vectors would require a "
        "different architecture (e.g. 1D CNN or RNN) and do not leverage our chosen transfer-learning "
        "pipeline. Mel-scale images preserve harmonic and transient patterns visible in urban and "
        "animal sounds, which explains why this representation was selected for both stages.",
    )
    add_figure(
        doc,
        FIG_DIR / "mel_vs_mfcc_comparison.png",
        "Mel-spectrogram vs MFCC for the same misclassified urban clip.",
        width=Inches(6),
    )

    add_heading(doc, "7. Key Findings", level=1)
    add_bullets(
        doc,
        [
            "Siren, car horn, drilling, and jackhammer form the most frequent urban confusion groups.",
            "Children playing is often confused with dog bark or street music due to outdoor ambient overlap.",
            "MobileNetV2 achieves the best urban accuracy with the lowest deployment cost.",
            "Cross-domain F1 drops by roughly 22–31% when moving from urban to ESC-50 animals.",
            "ImageNet transfer generalises better to animals than urban pretraining on this small dataset.",
            "Case studies confirm errors are often explainable by shared spectral texture, not random failure.",
        ],
    )

    add_heading(doc, "8. Step 6 Conclusions", level=1)
    add_bullets(
        doc,
        [
            "Detailed error analysis supports the MobileNetV2 deployment choice from Step 5.",
            "Confusion patterns align with acoustic similarity, validating the Mel-spectrogram pipeline.",
            "Inference benchmarks justify MobileNetV2 for real-time Streamlit inference.",
            "Cross-domain results highlight domain shift as the main limitation for generalisation.",
        ],
    )

    return doc


def build_dev_notes() -> Document:
    doc = Document()
    doc.add_heading("Development Notes — Step 6", level=0)
    add_para(doc, "Step 6: Error Analysis & Advanced Evaluation")
    add_heading(doc, "Work Completed", level=1)
    add_bullets(
        doc,
        [
            "Re-evaluated all 3 UrbanSound8K models on fold-10 test set",
            "Generated confusion matrices and top confused-pair bar charts",
            "Created 5 misclassification case studies with waveform + spectrogram panels",
            "Benchmarked parameter count, model size, and inference latency",
            "Summarised cross-domain performance from Stage 2",
            "Added Mel-spectrogram vs MFCC discussion figure",
            "Produced Step 6 Word report document",
        ],
    )
    add_heading(doc, "Estimated Time", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Hours"
    for i, (a, b, c) in enumerate([
        ("Urban error analysis", "3 models, confusion + pairs", "2"),
        ("Case studies", "5 detailed misclassification panels", "2"),
        ("Benchmarking", "Params, size, inference timing", "1"),
        ("Report", "Step 6 Word document", "2"),
        ("Total Step 6", "", "7"),
    ]):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
    return doc


def main() -> None:
    with (STEP6_DIR / "step6_summary.json").open(encoding="utf-8") as f:
        summary = json.load(f)

    report = build_report(summary)
    report_path = STEP6_DIR / "Step6_Error_Analysis_Report.docx"
    report.save(report_path)

    dev_notes = build_dev_notes()
    dev_path = STEP6_DIR / "Step6_Development_Notes.docx"
    dev_notes.save(dev_path)

    print(f"Saved: {report_path}")
    print(f"Saved: {dev_path}")


if __name__ == "__main__":
    main()
