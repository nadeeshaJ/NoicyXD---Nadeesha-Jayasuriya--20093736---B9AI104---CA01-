"""Generate the final combined assignment report, cover sheet, and development notes.

Run:
    python scripts/generate_final_report.py

Outputs:
    reports/final/Final_Assignment_Report.docx
    reports/final/Cover_Sheet.docx
    reports/final/Final_Development_Notes.docx
"""

from __future__ import annotations

import json
import platform
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import pandas as pd
import torch
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = PROJECT_ROOT / "reports" / "final"
FIG = PROJECT_ROOT / "reports" / "figures"


def load_json(path: Path) -> dict:
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def load_step_summaries() -> dict:
    return {
        "step1": load_json(PROJECT_ROOT / "reports" / "step1" / "step1_summary.json"),
        "step2": load_json(PROJECT_ROOT / "reports" / "step2" / "step2_summary.json"),
        "step3": load_json(PROJECT_ROOT / "reports" / "step3" / "step3_summary.json"),
        "step4": load_json(PROJECT_ROOT / "reports" / "step4" / "step4_summary.json"),
        "step5": load_json(PROJECT_ROOT / "reports" / "step5" / "step5_verification.json"),
        "step6": load_json(PROJECT_ROOT / "reports" / "step6" / "step6_summary.json"),
    }


def load_config() -> dict:
    with (PROJECT_ROOT / "config" / "config.yaml").open(encoding="utf-8") as f:
        return yaml.safe_load(f)


def get_env_info() -> dict:
    return {
        "python": platform.python_version(),
        "platform": platform.platform(),
        "torch": torch.__version__,
        "cuda_available": torch.cuda.is_available(),
        "cuda_device": torch.cuda.get_device_name(0) if torch.cuda.is_available() else "N/A",
    }


def load_ablation_summary() -> dict | None:
    path = FINAL_DIR / "ca1_ablation_summary.json"
    if path.exists():
        return load_json(path)
    return None


def load_urban_test_metrics(model_name: str) -> dict:
    path = PROJECT_ROOT / "experiments" / "urbansound8k" / model_name / "test_metrics.json"
    return load_json(path)


def mobilenetv2_per_class_rows() -> list[list[str]]:
    """Per-class precision, recall, F1 for best urban model (fold-10 test)."""
    report = load_urban_test_metrics("mobilenetv2")["classification_report"]
    rows = []
    for cls in sorted(k for k in report if k not in ("accuracy", "macro avg", "weighted avg")):
        row = report[cls]
        rows.append([
            cls.replace("_", " "),
            f"{row['precision']:.3f}",
            f"{row['recall']:.3f}",
            f"{row['f1-score']:.3f}",
            str(int(row["support"])),
        ])
    macro = report["macro avg"]
    rows.append([
        "Macro average",
        f"{macro['precision']:.3f}",
        f"{macro['recall']:.3f}",
        f"{macro['f1-score']:.3f}",
        str(int(macro["support"])),
    ])
    return rows


def urban_model_metrics_rows() -> list[list[str]]:
    rows = []
    for model_name, ca1_label in (
        ("custom_cnn", "Model 1 — Custom CNN (baseline)"),
        ("resnet50", "Model 2 — ResNet50 (transfer)"),
        ("mobilenetv2", "Model 3 — MobileNetV2 (custom head + fine-tune)"),
    ):
        metrics = load_urban_test_metrics(model_name)
        summary_path = PROJECT_ROOT / "experiments" / "urbansound8k" / model_name / "training_summary.json"
        train_time = load_json(summary_path).get("train_time_sec", 0)
        macro = metrics["classification_report"]["macro avg"]
        rows.append([
            ca1_label,
            f"{metrics['metrics']['accuracy']:.4f}",
            f"{macro['recall']:.4f}",
            f"{metrics['metrics']['macro_f1']:.4f}",
            f"{train_time:.0f}s",
        ])
    return rows


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_figure(doc: Document, path: Path, caption: str, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)
    doc.add_paragraph()


def build_cover_sheet() -> Document:
    doc = Document()
    doc.add_heading("Dublin Business School", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Assignment Cover Sheet", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Details"],
        [
            ["Student Name", "Nadeesha Jayasuriya"],
            ["Student Number", "[Your Student ID]"],
            ["Programme", "[Your Programme]"],
            ["Module", "B9AI104 Deep Learning"],
            ["Lecturer", "Dr. Shahram Azizi Sazi"],
            ["Assessment", "Continuous Assessment 1 (CA1) — 40%"],
            [
                "Assignment Title",
                "Deep Learning-Based Environmental Sound Classification Using "
                "Mel-Spectrogram Images: A Cross-Domain Study of Urban and Animal Sounds",
            ],
            ["Submission Date", date.today().strftime("%d %B %Y")],
            ["Estimated Time to Complete", "65 hours (see Development Notes document)"],
        ],
    )
    add_para(doc, "Declaration", bold=True)
    add_para(
        doc,
        "I confirm that this submission is my own work. All sources, datasets, and references "
        "have been acknowledged in the bibliography. I understand that plagiarism or improper "
        "use of generative AI without declaration may result in disciplinary action.",
    )
    add_para(doc, "Generative AI Use Statement", bold=True)
    add_para(
        doc,
        "Generative AI tools (e.g. Cursor AI) were used for code scaffolding, debugging, and "
        "report formatting assistance. All modelling decisions, experiments, results, analysis, "
        "and written interpretation were reviewed and validated by me.",
    )
    return doc


def add_cover_page_to_report(doc: Document) -> None:
    """Formal DBS cover sheet as the first page of the main report."""
    doc.add_heading("Dublin Business School", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Assignment Cover Sheet", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Details"],
        [
            ["Student Name", "Nadeesha Jayasuriya"],
            ["Student Number", "[Your Student ID]"],
            ["Programme", "[Your Programme]"],
            ["Module", "B9AI104 Deep Learning"],
            ["Lecturer", "Dr. Shahram Azizi Sazi"],
            ["Assessment", "Continuous Assessment 1 (CA1) — 40%"],
            [
                "Assignment Title",
                "Deep Learning-Based Environmental Sound Classification Using "
                "Mel-Spectrogram Images: A Cross-Domain Study of Urban and Animal Sounds",
            ],
            ["Submission Date", date.today().strftime("%d %B %Y")],
        ],
    )
    add_para(
        doc,
        "Declaration: I confirm this submission is my own work. Sources are cited in the "
        "bibliography. AI tools were used for coding/debugging support only; all experiments "
        "and interpretations were verified by me.",
    )
    doc.add_page_break()


def build_main_report(data: dict, cfg: dict, env: dict, ablation: dict | None = None) -> Document:
    s1, s2, s3, s4, s5, s6 = (
        data["step1"],
        data["step2"],
        data["step3"],
        data["step4"],
        data["step5"],
        data["step6"],
    )
    urban = s1["urbansound8k"]
    esc = s1["esc50_animals"]
    ablation = ablation or {}
    head_ablation = ablation.get("model3_head_ablation", [])
    dropout_ablation = ablation.get("custom_cnn_dropout_ablation", [])
    urban_classes = ", ".join(cfg["datasets"]["urbansound8k"]["classes"])
    animal_classes = ", ".join(cfg["datasets"]["esc50_animals"]["classes"])

    doc = Document()
    add_cover_page_to_report(doc)
    title = doc.add_heading(
        "Deep Learning-Based Environmental Sound Classification Using Mel-Spectrogram Images",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "B9AI104 Deep Learning — Final Assignment Report")
    doc.add_paragraph()

    # --- 1. Introduction ---
    add_heading(doc, "1. Introduction", level=1)
    add_heading(doc, "1.0 CA1 Brief Introduction (Required Format)", level=2)
    add_para(
        doc,
        "I am doing multi-class environmental sound classification on unstructured Mel-spectrogram "
        "RGB images derived from audio. The application is urban noise monitoring and animal sound "
        "recognition for smart-city and wildlife sensing. My data consists of 8,732 UrbanSound8K "
        "clips (10 urban classes, fold-10 test) and 200 ESC-50 animal clips (10 classes, Hugging Face subset).",
        bold=True,
    )
    add_heading(doc, "1.0.1 CRISP-DM Report Mapping", level=2)
    add_table(
        doc,
        ["CRISP-DM Phase", "Report Section", "Content"],
        [
            ["Business Understanding", "Section 1", "Problem, application, research questions"],
            ["Data Understanding", "Section 3", "EDA, class distribution, dataset counts"],
            ["Data Preparation", "Section 3.3", "Normalisation, Mel-spec pipeline, SpecAugment"],
            ["Modeling", "Sections 2 & 4", "Three CNN architectures, training strategy"],
            ["Evaluation", "Section 5", "Accuracy, recall, confusion matrix, probabilities"],
        ],
    )
    add_para(
        doc,
        "Environmental sound classification supports applications in urban monitoring, "
        "wildlife detection, and smart-city analytics. This project converts audio clips "
        "into Mel-spectrogram RGB images and classifies them using convolutional neural networks. "
        "A two-stage experimental design evaluates performance on UrbanSound8K (10 urban classes) "
        "and extends the same pipeline to an ESC-50 animal subset (10 classes) for cross-domain analysis.",
    )
    add_heading(doc, "1.0.2 Class Names (Presentation Requirement)", level=2)
    add_para(doc, "UrbanSound8K classes (10):", bold=True)
    add_para(doc, urban_classes.replace("_", " "))
    add_para(doc, "ESC-50 animal classes (10):", bold=True)
    add_para(doc, animal_classes)
    add_heading(doc, "1.1 Problem Statement & Research Questions", level=2)
    add_bullets(
        doc,
        [
            "Can Mel-spectrogram images classified by CNNs achieve strong performance on urban environmental sounds?",
            "How do a custom CNN, ResNet50, and MobileNetV2 compare on the same preprocessed data?",
            "Does the preprocessing pipeline generalise from urban to animal sounds?",
            "Which transfer-learning strategy works best on a small ESC-50 subset?",
            "Which model is most suitable for real-time deployment?",
        ],
    )
    add_heading(doc, "1.2 Contributions", level=2)
    add_bullets(
        doc,
        [
            "End-to-end reproducible pipeline: audio → Mel-spec → 224×224 RGB → CNN → prediction",
            "Three-model comparison on UrbanSound8K with official fold-10 test evaluation",
            "Cross-domain ESC-50 extension comparing urban transfer, scratch training, and ImageNet transfer",
            "Detailed error analysis with case studies and inference benchmarking",
            "Streamlit web application for live demo of both urban and animal modes",
        ],
    )

    doc.add_page_break()

    # --- 2. Model Architecture Background ---
    add_heading(doc, "2. Model Architecture Background", level=1)
    add_heading(doc, "2.1 Audio-to-Image Representation", level=2)
    add_para(
        doc,
        "Raw waveforms are transformed using the Short-Time Fourier Transform (STFT), "
        "projected onto a Mel filterbank to match human frequency perception, converted to "
        "decibel scale, min-max normalised, and resized to 224×224 RGB images suitable for CNN input.",
    )
    add_bullets(
        doc,
        [
            "STFT: n_fft=2048, hop_length=512",
            "Mel filterbank: 128 bins, fmin=0, fmax=11025 Hz",
            "Power spectrogram converted to dB: 10·log10(power / ref)",
            "Normalisation: min-max to [0, 1] per clip; replicated to 3 channels",
        ],
    )
    add_figure(
        doc,
        FIG / "step2" / "preprocessing_pipeline_diagram.png",
        "Figure 2.1 — Preprocessing pipeline from audio to CNN input.",
    )

    add_heading(doc, "2.2 Custom CNN — Conventional Baseline (Model 1)", level=2)
    add_para(
        doc,
        "Model 1 is a conventional CNN baseline built from standard blocks: convolution, ReLU, "
        "max pooling, flatten, dense, dropout, and softmax. It is trained from scratch (not the "
        "CA1 customised transfer-learning model — that role belongs to Model 3).",
    )
    add_bullets(
        doc,
        [
            "Four convolutional blocks (32→64→128→128 filters) with ReLU and max pooling",
            "Fully connected head: 256 units, dropout 0.5, 10-class softmax output",
            "Trained from scratch without pretrained weights",
            f"Parameters: {s3['model_specs'][0]['total_parameters']:,}",
        ],
    )
    add_figure(
        doc,
        FIG / "step3" / "architecture_custom_cnn.png",
        "Figure 2.2 — Custom CNN architecture diagram.",
        width=Inches(4.5),
    )
    add_figure(
        doc,
        FIG / "final" / "model_summary_custom_cnn.png",
        "Figure 2.3 — model.summary() parameter count (Custom CNN, 224×224×3 input).",
        width=Inches(6.2),
    )

    add_heading(doc, "2.3 ResNet50 — Traditional Transfer Learning (Model 2)", level=2)
    add_bullets(
        doc,
        [
            "ImageNet-pretrained ResNet50 backbone with residual skip connections",
            "Two-phase training: frozen backbone + head training, then fine-tuning",
            "ImageNet mean/std normalisation applied at inference",
            f"Parameters: {s3['model_specs'][1]['total_parameters']:,}",
        ],
    )
    add_figure(
        doc,
        FIG / "final" / "model_summary_resnet50.png",
        "Figure 2.4 — model.summary() parameter count (ResNet50 + replaced 10-class head).",
        width=Inches(6.2),
    )

    add_heading(doc, "2.4 MobileNetV2 — Customised Transfer Learning (Model 3)", level=2)
    add_bullets(
        doc,
        [
            "Depthwise separable convolutions for parameter efficiency",
            "Same transfer-learning strategy as ResNet50",
            "Customised classifier head (dropout + linear) with head-size ablation — CA1 Model 3",
            f"Parameters: {s3['model_specs'][2]['total_parameters']:,}",
        ],
    )
    add_figure(
        doc,
        FIG / "step3" / "architecture_transfer_learning.png",
        "Figure 2.5 — Transfer learning strategy for ResNet50 and MobileNetV2.",
    )
    add_figure(
        doc,
        FIG / "final" / "model_summary_mobilenetv2.png",
        "Figure 2.6 — model.summary() parameter count (MobileNetV2 + custom 10-class head).",
        width=Inches(6.2),
    )

    add_heading(doc, "2.5 CNN Concepts (Examiner Q&A)", level=2)
    add_para(
        doc,
        "Weight sharing: Unlike a fully connected ANN where each neuron has unique weights, a CNN "
        "reuses the same filter kernel across the entire image. This provides translation invariance "
        "and dramatically reduces trainable parameters.",
    )
    add_para(
        doc,
        "ML vs DL: Deep learning models contain hidden convolutional and dense layers between input "
        "and output; logistic regression maps input directly to output with no hidden layer.",
    )
    add_para(
        doc,
        "Back propagation: After feed-forward computes predicted class probabilities, cross-entropy loss "
        "measures error between predicted and actual labels. Back propagation calculates the gradient of "
        "loss with respect to each weight, and Adam updates weights (new weight = old weight − learning "
        "rate × gradient). This repeats each epoch until validation loss stops improving.",
    )
    add_para(
        doc,
        "Architecture boundary: Convolution + pooling layers (before Flatten) extract spatial features; "
        "Flatten → Dense → Dropout → Softmax is the ANN classification head. Transfer-learning diagrams "
        "mark the frozen ImageNet backbone separately from the trainable custom head.",
    )

    add_heading(doc, "2.6 Mathematical Foundations", level=2)
    add_para(
        doc,
        "STFT and Mel-spectrogram: The short-time Fourier transform computes local frequency content. "
        "For a frame x[n], the STFT is X(k, m) = Σ_n x[n] w[n−mH] e^(−j2πkn/N), where w is the window, "
        "H is hop length (512), and N is n_fft (2048). Mel filterbank energies are mapped to "
        "perceptual pitch using m = 2595·log10(1 + f/700), then converted to decibels: "
        "S_dB = 10·log10(S / S_ref). Only magnitude is retained — phase is discarded (see Section 7.3).",
    )
    add_para(
        doc,
        "ResNet50 residual block: To ease optimisation of deep networks, residual mapping defines "
        "H(x) = F(x) + x, where x is the input to the block and F(x) is the learned residual "
        "transformation (convolution stack). Skip connections allow gradients to flow directly and "
        "reduce vanishing-gradient problems in very deep CNNs.",
    )
    add_para(
        doc,
        "MobileNetV2 depthwise separable convolution: Standard convolution cost is "
        "D_K × D_K × M × N × H × W. Depthwise separable convolution splits this into (1) depthwise "
        "convolution (D_K × D_K × M × H × W) and (2) pointwise 1×1 convolution (M × N × H × W), "
        "reducing computation and parameters — critical for efficient deployment on limited hardware.",
    )
    add_para(
        doc,
        "Classification output: Hidden layers use ReLU; the final layer uses softmax so class "
        "probabilities sum to 1: P(class=k) = exp(z_k) / Σ_j exp(z_j). Cross-entropy loss for "
        "true class y is L = −log P(y).",
    )

    add_heading(doc, "2.7 Model Rules (Operational Constraints)", level=2)
    add_para(doc, "Custom CNN (Model 1) — rules:", bold=True)
    add_bullets(
        doc,
        [
            "Input: fixed 224×224×3 Mel-spec RGB; pixel values in [0, 1] after normalisation",
            "Layer order: Conv2d → ReLU → MaxPool (×4 blocks) → Flatten → Dense(256) → ReLU → Dropout(0.5) → Dense(10)",
            "Output: 10 logits; softmax applied at inference; no ImageNet normalisation",
            "Training: all layers trainable from epoch 1; no freeze/unfreeze phases",
            "Loss: cross-entropy with integer labels; optimiser Adam, lr=1e-4",
        ],
    )
    add_para(doc, "ResNet50 (Model 2) — rules:", bold=True)
    add_bullets(
        doc,
        [
            "Input: 224×224×3 with ImageNet mean/std normalisation (μ=[0.485,0.456,0.406], σ=[0.229,0.224,0.225])",
            "Backbone: pretrained ImageNet conv blocks with residual skip connections",
            "Phase 1: freeze all backbone layers; train only fc head (requires_grad on fc only)",
            "Phase 2: unfreeze top residual blocks + fc; backbone lr=1e-5, head lr=1e-4",
            "Output: replaced fc head with 10-class softmax (not ImageNet 1000 classes)",
            "Checkpoint rule: save model with lowest validation loss, not training loss",
        ],
    )
    add_para(doc, "MobileNetV2 (Model 3) — rules:", bold=True)
    add_bullets(
        doc,
        [
            "Input: same 224×224×3 ImageNet normalisation as ResNet50",
            "Backbone: inverted residual blocks with depthwise separable convolutions",
            "Phase 1: freeze features; train classifier head only (~3 epochs)",
            "Phase 2: unfreeze last 4 inverted residual blocks + classifier; differential learning rates",
            "Custom head: Dropout(0.2) → Linear(1280 → 10); ablation tested 64/128/256 hidden nodes",
            "Deployment rule: selected as best model — highest macro F1 with smallest checkpoint size "
            "(not lowest inference latency; Custom CNN is faster but less accurate)",
        ],
    )

    doc.add_page_break()

    # --- 3. Dataset Implementation ---
    add_heading(doc, "3. Dataset Implementation", level=1)
    add_heading(doc, "3.1 UrbanSound8K (Stage 1)", level=2)
    add_table(
        doc,
        ["Property", "Value"],
        [
            ["Total clips", str(urban["total_clips"])],
            ["Classes", str(urban["num_classes"])],
            ["Duration", f"{urban['duration_min_sec']}–{urban['duration_max_sec']} s (mean {urban['duration_mean_sec']:.2f} s)"],
            ["Folds", f"10 (test = fold {urban['test_fold']}, train = folds {urban['train_folds']})"],
            ["Train / val / test", "7105 / 790 / 837 processed clips"],
        ],
    )
    add_figure(
        doc,
        FIG / "step1" / "urban_class_distribution.png",
        "Figure 3.1 — UrbanSound8K class distribution.",
        width=Inches(5),
    )

    add_heading(doc, "3.2 ESC-50 Animals (Stage 2)", level=2)
    add_table(
        doc,
        ["Property", "Value"],
        [
            ["Total clips", str(esc["total_clips"])],
            ["Classes", str(esc["num_classes"])],
            ["Duration", "5 s (trimmed/padded to 4 s for pipeline consistency)"],
            ["Split", "70% train / 15% val / 15% test (stratified)"],
            ["Train / val / test", "140 / 30 / 30 processed clips"],
        ],
    )
    add_figure(
        doc,
        FIG / "step1" / "esc50_class_distribution.png",
        "Figure 3.2 — ESC-50 animal subset class distribution.",
        width=Inches(5),
    )

    add_heading(doc, "3.3 Preprocessing Pipeline", level=2)
    add_bullets(
        doc,
        [
            f"Resample to {cfg['audio']['sample_rate']} Hz mono; pad/trim to {cfg['audio']['duration_sec']} s",
            f"Mel-spectrogram: {cfg['spectrogram']['n_mels']} bins, saved as {cfg['image']['height']}×{cfg['image']['width']} PNG",
            "SpecAugment during training: time/frequency masking",
            "Zero preprocessing errors across all splits (Step 2 validation)",
        ],
    )
    add_para(
        doc,
        "Sample rate (22,050 Hz): This is a standard librosa default that preserves the Nyquist "
        "frequency of 11,025 Hz — sufficient for urban and animal environmental sounds while halving "
        "storage and compute versus the original 44.1 kHz recordings. It balances frequency coverage "
        "with training efficiency on CPU/GPU.",
    )
    add_para(
        doc,
        "Mel bins (128): 128 bins provide fine enough frequency resolution to separate classes such as "
        "siren vs car horn or drilling vs jackhammer, without producing excessively large spectrogram "
        "matrices. Combined with 224×224 resizing, this yields a compact image representation suitable "
        "for ImageNet-scale CNN inputs.",
    )
    add_figure(
        doc,
        FIG / "step2" / "preprocessing_examples_urbansound8k.png",
        "Figure 3.3 — Preprocessing examples (waveform → Mel-spec → RGB).",
        width=Inches(6),
    )

    doc.add_page_break()

    # --- 4. Deep Learning Design ---
    add_heading(doc, "4. Deep Learning Design", level=1)
    add_heading(doc, "4.1 Design Choices", level=2)
    add_table(
        doc,
        ["Setting", "Value", "Rationale"],
        [
            ["Input size", "224×224×3", "Standard ImageNet input; enables transfer learning"],
            ["Batch size", str(cfg["training"]["batch_size"]), "Balance GPU memory and gradient stability"],
            ["Learning rate", str(cfg["training"]["learning_rate"]), "Conservative LR for fine-tuning"],
            ["Optimiser", "Adam", "Standard choice for CNN audio-image tasks"],
            ["Loss", "Cross-entropy", "Multi-class classification"],
            ["Early stopping", f"patience={cfg['training']['early_stopping_patience']}", "Prevent overfitting"],
            ["Random seed", str(cfg["training"]["seed"]), "Reproducibility"],
        ],
    )

    add_heading(doc, "4.2 Training Strategy", level=2)
    add_bullets(
        doc,
        [
            "Custom CNN: full training from random initialisation",
            "ResNet50 / MobileNetV2: Phase 1 — train classifier head with frozen backbone; Phase 2 — fine-tune",
            "Stage 1 trained on CPU; Stage 2 and deployment inference used GPU (RTX 4060 Laptop)",
            "Best checkpoint selected by validation loss",
        ],
    )

    add_heading(doc, "4.3 CA1 Three-Model Structure", level=2)
    add_table(
        doc,
        ["CA1 Role", "Model", "Description"],
        [
            [
                "Model 1",
                "Custom CNN",
                "Conventional CNN baseline from scratch (conv, pooling, flatten, dense, dropout, softmax)",
            ],
            ["Model 2", "ResNet50", "Traditional transfer learning — ImageNet backbone + replaced 10-class head"],
            [
                "Model 3",
                "MobileNetV2",
                "Customised transfer model — frozen backbone + tuned classifier head + fine-tune",
            ],
        ],
    )
    add_para(
        doc,
        "Per CA1 guidelines: Model 1 is the conventional from-scratch baseline; Models 2 and 3 are "
        "traditional transfer-learning CNNs, with Model 3 as the customised head + fine-tune experiment. "
        "Model 3 (MobileNetV2) achieved the highest urban accuracy and recall and was selected for deployment. "
        "Hyperparameter experiments on the custom head (Section 4.4) followed the course guideline of testing "
        "64, 128, and 256 hidden nodes with a frozen ImageNet backbone.",
    )

    add_heading(doc, "4.4 Hyperparameter Experiments (CA1 Requirement)", level=2)
    if head_ablation:
        head_rows = []
        for row in head_ablation:
            head_rows.append([
                str(row.get("hidden_units", "")),
                f"{row.get('test_accuracy', 0):.4f}",
                f"{row.get('test_macro_recall', 0):.4f}",
                f"{row.get('test_macro_f1', 0):.4f}",
                f"{row.get('train_time_sec', 0):.0f}s",
            ])
        add_table(
            doc,
            ["Head hidden units", "Test accuracy", "Macro recall", "Macro F1", "Train time"],
            head_rows,
        )
        add_figure(
            doc,
            FIG / "final" / "ablation_model3_head_sizes.png",
            "Figure 4.1 — Model 3 head-size ablation (MobileNetV2, frozen backbone, 8 epochs).",
            width=Inches(5.5),
        )
        best_head = max(head_ablation, key=lambda r: r.get("test_macro_f1", 0))
        add_para(
            doc,
            f"Best head configuration in ablation: {best_head.get('head_config')} "
            f"(macro F1={best_head.get('test_macro_f1', 0):.4f}). "
            "The full Model 3 production run used extended fine-tuning and early stopping, "
            "achieving higher final test performance than these short head-only trials.",
        )
    else:
        add_para(
            doc,
            "Run scripts/run_ca1_ablation_studies.py to generate Model 3 head-size comparison results.",
        )

    add_heading(doc, "4.5 Dropout Comparison (Custom CNN)", level=2)
    if dropout_ablation:
        drop_rows = []
        for row in dropout_ablation:
            drop_rows.append([
                str(row.get("dropout", "")),
                f"{row.get('test_accuracy', 0):.4f}",
                f"{row.get('test_macro_recall', 0):.4f}",
                f"{row.get('train_time_sec', 0):.0f}s",
            ])
        add_table(doc, ["Dropout rate", "Test accuracy", "Macro recall", "Train time"], drop_rows)
        add_figure(
            doc,
            FIG / "final" / "ablation_custom_cnn_dropout.png",
            "Figure 4.2 — Custom CNN dropout 0.0 vs 0.5 (12-epoch ablation).",
            width=Inches(4.5),
        )
        add_para(
            doc,
            "Dropout 0.5 was used in the final Custom CNN because it reduces overfitting by randomly "
            "dropping 50% of connections during training, following course examples (AlexNet/VGGNet).",
        )
    else:
        add_para(doc, "Final Custom CNN uses dropout=0.5 in the classifier head (see Section 2.2).")

    add_heading(doc, "4.6 Strengths and Weaknesses", level=2)
    add_bullets(
        doc,
        [
            "Strength: Unified Mel-image pipeline works for both urban and animal domains without modification",
            "Strength: Transfer learning substantially outperforms custom CNN on limited data",
            "Weakness: Fixed 4 s window may truncate or pad clips awkwardly",
            "Weakness: Mel-spec loses phase information; similar-sounding classes remain confusable",
            "Weakness: ESC-50 subset is very small (200 clips), limiting cross-domain conclusions",
        ],
    )

    add_heading(doc, "4.7 Early Stopping and Augmentation Justification", level=2)
    add_bullets(
        doc,
        [
            f"EarlyStopping equivalent: training monitors validation loss with patience={cfg['training']['early_stopping_patience']}; "
            "best checkpoint restored when validation stops improving.",
            "SpecAugment applied on training Mel-spec images only — never on validation or test splits.",
            "Augmentation justified: urban classes vary in recording conditions; time/frequency masking "
            "increases effective training diversity without leaking to test data.",
        ],
    )

    doc.add_page_break()

    # --- 5. Evaluation ---
    add_heading(doc, "5. Model Comparison, Evaluation & Validation", level=1)
    add_heading(doc, "5.1 UrbanSound8K Results — Accuracy & Recall (Fold 10 Test)", level=2)
    add_para(
        doc,
        "All three CA1 models evaluated on the official fold-10 test set. Macro recall averages "
        "recall across all 10 classes. Per-class metrics for the best model are in Table 5.2.",
    )
    add_table(
        doc,
        ["Model (CA1 role)", "Accuracy", "Macro recall", "Macro F1", "Training time"],
        urban_model_metrics_rows(),
    )

    comparison_path = PROJECT_ROOT / "reports" / "step3" / "model_comparison_urbansound8k.csv"
    if comparison_path.exists():
        comp_df = pd.read_csv(comparison_path)
        rows = []
        for _, row in comp_df.iterrows():
            m = load_urban_test_metrics(row["model"])
            rows.append([
                row.get("model", ""),
                f"{row.get('accuracy', 0):.4f}",
                f"{m['classification_report']['macro avg']['recall']:.4f}",
                f"{row.get('macro_f1', 0):.4f}",
            ])
        add_table(doc, ["Model", "Accuracy", "Macro recall", "Macro F1"], rows)

    add_para(doc, f"Best urban model: {s3['best_model']} (macro F1 = {s3['best_macro_f1']:.4f})", bold=True)

    add_heading(doc, "5.1.1 Per-Class Classification Report — MobileNetV2 (Best Model)", level=2)
    add_para(
        doc,
        "Table 5.2 reports precision, recall, and F1 per class on the fold-10 test set. Recall is "
        "reported explicitly as required by the CA1 evaluation checklist.",
    )
    add_table(
        doc,
        ["Class", "Precision", "Recall", "F1-score", "Support"],
        mobilenetv2_per_class_rows(),
    )

    add_figure(
        doc,
        FIG / "step3" / "model_comparison_urbansound8k.png",
        "Figure 5.1 — UrbanSound8K model comparison.",
    )

    add_heading(doc, "5.1.2 Training Curves & Best Epoch Selection", level=2)
    add_para(
        doc,
        "Training and validation loss/accuracy curves were saved for all models. The best checkpoint "
        "was selected by minimum validation loss (not training loss), with early stopping when "
        "validation loss failed to improve for 5 consecutive epochs.",
    )
    for fig_num, (model_name, caption) in enumerate(
        (
            ("custom_cnn", "Custom CNN (Model 1)"),
            ("resnet50", "ResNet50 (Model 2)"),
            ("mobilenetv2", "MobileNetV2 (Model 3)"),
        ),
        start=2,
    ):
        add_figure(
            doc,
            FIG / "step3" / "urbansound8k" / model_name / "training_history.png",
            f"Figure 5.{fig_num} — Training curves: {caption}.",
            width=Inches(5.2),
        )

    add_heading(doc, "5.1.3 Probabilities & Explainable Predictions", level=2)
    add_para(
        doc,
        "The Streamlit deployment reports softmax probabilities for all classes. Example urban inference: "
        f"dog_bark predicted with {s5['checks'][0]['confidence']:.1%} confidence "
        f"(not just the class label). Top-3 predictions with confidence bars are shown to support explainable AI.",
    )

    add_heading(doc, "5.2 ESC-50 Transfer Learning Results", level=2)
    esc_rows = []
    for run in s4["esc50_runs"]:
        esc_rows.append([
            run["run_name"],
            f"{run['test_metrics']['accuracy']:.4f}",
            f"{run['test_metrics']['macro_f1']:.4f}",
        ])
    add_table(doc, ["Approach", "Accuracy", "Macro F1"], esc_rows)
    add_para(
        doc,
        f"Best ESC-50 approach: {s4['best_esc50_run']} (macro F1 = {s4['best_esc50_macro_f1']:.4f})",
        bold=True,
    )

    add_heading(doc, "5.3 Cross-Domain Comparison", level=2)
    cross_rows = []
    for row in s4["cross_domain"]:
        cross_rows.append([
            row["approach"],
            row["domain"],
            f"{row['macro_f1']:.4f}",
            f"{row['f1_drop']:.4f}",
        ])
    add_table(doc, ["Approach", "Domain", "Macro F1", "F1 Drop vs Urban"], cross_rows)
    add_figure(
        doc,
        FIG / "step4" / "cross_domain_comparison.png",
        "Figure 5.5 — Cross-domain macro F1 comparison.",
    )
    add_para(
        doc,
        "ImageNet-only MobileNetV2 generalised best to animal sounds. UrbanSound8K transfer "
        "did not outperform ImageNet pretraining on this small dataset, indicating significant domain shift.",
    )

    add_heading(doc, "5.4 Error Analysis", level=2)
    best_urban_err = next(
        r for r in s6["urban_error_analysis"] if r["run_name"] == s6["best_urban_run"]
    )
    add_para(doc, f"Top confused pairs ({s6['best_urban_run']}):", bold=True)
    for pair in best_urban_err["top_confused_pairs"][:5]:
        add_bullets(
            doc,
            [f"{pair['true_class']} → {pair['predicted_class']} ({pair['count']} times)"],
        )
    add_figure(
        doc,
        FIG / "step6" / f"confusion_matrix_{s6['best_urban_run']}.png",
        "Figure 5.6 — Confusion matrix (MobileNetV2, urban test set).",
        width=Inches(4.8),
    )
    add_figure(
        doc,
        FIG / "step6" / "case_studies" / "case_study_01.png",
        "Figure 5.7 — Misclassification case study (siren predicted as children playing).",
        width=Inches(6),
    )

    add_heading(doc, "5.5 Inference Benchmarking", level=2)
    add_para(
        doc,
        "Inference was measured on GPU (100 runs, batch size 1). Custom CNN has the lowest latency "
        "(~0.91 ms) but lower accuracy. MobileNetV2 was selected for deployment because it achieved "
        "the highest macro F1 while keeping the smallest checkpoint — a better overall "
        "accuracy-efficiency trade-off than latency alone.",
    )
    bench_rows = []
    for b in s6["inference_benchmarks"]:
        bench_rows.append([
            b["model"],
            f"{b['total_parameters']:,}",
            f"{b['model_file_size_mb']:.2f} MB",
            f"{b['inference_ms_mean']:.2f} ms",
            f"{b['test_macro_f1']:.4f}",
        ])
    add_table(
        doc,
        ["Model", "Parameters", "File Size", "Inference (GPU)", "Macro F1"],
        bench_rows,
    )
    add_figure(
        doc,
        FIG / "step6" / "mel_vs_mfcc_comparison.png",
        "Figure 5.8 — Mel-spectrogram vs MFCC (why Mel images suit CNNs).",
        width=Inches(6),
    )

    doc.add_page_break()

    # --- 6. Deployment ---
    add_heading(doc, "6. Application Deployment", level=1)
    add_para(
        doc,
        "A Streamlit web application (`app/streamlit_app.py`) deploys the trained classifiers "
        "for live inference. Users select Urban or Animal mode, upload a WAV file, and receive "
        "waveform, Mel-spectrogram, model input preview, and top-3 predictions.",
    )
    add_table(
        doc,
        ["Mode", "Model", "Checkpoint"],
        [
            ["Urban", s5["deployment"]["urban"]["model_name"], s5["deployment"]["urban"]["checkpoint"]],
            ["Animal", s5["deployment"]["animal"]["model_name"], s5["deployment"]["animal"]["checkpoint"]],
        ],
    )
    add_bullets(
        doc,
        [
            f"Urban demo: {Path(s5['checks'][0]['sample_audio']).name} → {s5['checks'][0]['prediction']} ({s5['checks'][0]['confidence']:.1%})",
            f"Animal demo: {Path(s5['checks'][1]['sample_audio']).name} → {s5['checks'][1]['prediction']} ({s5['checks'][1]['confidence']:.1%})",
            "Run command: python -m streamlit run app/streamlit_app.py",
        ],
    )
    add_figure(
        doc,
        FIG / "step5" / "app_demo_urban.png",
        "Figure 6.1 — Streamlit urban mode demo preview.",
        width=Inches(6),
    )
    add_figure(
        doc,
        FIG / "step5" / "app_demo_animal.png",
        "Figure 6.2 — Streamlit animal mode demo preview.",
        width=Inches(6),
    )

    doc.add_page_break()

    # --- 7. Results & Elaboration ---
    add_heading(doc, "7. Results & Elaboration", level=1)
    add_heading(doc, "7.1 Best Model Selection", level=2)
    add_para(
        doc,
        "MobileNetV2 was selected for deployment because it achieved the highest UrbanSound8K "
        "macro F1 (0.831) while keeping the smallest checkpoint (8.76 MB). Although the Custom CNN "
        "had lower GPU inference latency (~0.91 ms vs ~4.2 ms), MobileNetV2 provided a better overall "
        "accuracy-efficiency trade-off and serves as the backbone for both Streamlit modes.",
    )
    add_heading(doc, "7.2 Key Implications", level=2)
    add_bullets(
        doc,
        [
            "Mel-spectrogram images enable reuse of proven CNN architectures for audio classification",
            "Transfer learning is essential when training data is limited (ESC-50: 140 train clips)",
            "Domain shift between urban and animal sounds remains the primary generalisation challenge",
            "Confusion errors are often acoustically explainable (siren/horn, drilling/jackhammer)",
        ],
    )
    add_para(
        doc,
        "Link to spectrogram limitations: Many confusions (e.g. siren → children_playing, drilling → "
        "jackhammer) occur because the Mel-spectrogram uses magnitude-only power (phase discarded during "
        "STFT). Classes with overlapping magnitude patterns but different temporal phase structure can "
        "appear nearly identical as 224×224 images — explaining why errors cluster between acoustically "
        "similar urban classes rather than being random mislabels.",
    )
    add_heading(doc, "7.3 Limitations", level=2)
    add_bullets(
        doc,
        [
            "Stage 1 was trained on CPU (~4.7 hours total) — a hardware constraint mitigated by choosing "
            "MobileNetV2 (efficient depthwise separable convolutions, 8.8 MB checkpoint) as the deployment model",
            "Magnitude spectrograms discard phase information, contributing to confusions between similar classes",
            "ESC-50 animal subset has only 30 test clips — high variance in metrics",
            "No real-time audio recording in the app (upload only)",
            "Single fixed preprocessing config; no per-class tuning",
        ],
    )
    add_heading(doc, "7.4 Future Work", level=2)
    add_bullets(
        doc,
        [
            "Audio augmentation beyond SpecAugment (time stretch, pitch shift)",
            "Attention-based or transformer models on spectrogram patches",
            "Larger animal dataset and multi-domain fine-tuning",
            "Model quantisation for edge deployment",
            "Ensemble of urban + animal experts with domain detection",
        ],
    )

    doc.add_page_break()

    # --- 8. Conclusion ---
    add_heading(doc, "8. Conclusion", level=1)
    add_para(
        doc,
        "This project implemented a complete deep learning pipeline for environmental sound "
        "classification using Mel-spectrogram images. On UrbanSound8K, MobileNetV2 achieved "
        "82.7% accuracy and 0.831 macro F1, outperforming a custom CNN (75.0%) and ResNet50 (81.2%). "
        "Cross-domain extension to ESC-50 animals confirmed that the preprocessing pipeline "
        "generalises, but performance drops significantly due to domain shift and limited data. "
        "ImageNet transfer learning proved more effective than urban pretraining on the small "
        "animal subset. Detailed error analysis, inference benchmarks, and a Streamlit deployment "
        "demonstrate both the strengths and limitations of the Mel-spectrogram CNN approach.",
    )

    doc.add_page_break()

    # --- 9. Bibliography ---
    add_heading(doc, "9. Bibliography", level=1)
    references = [
        "Salamon, J. and Bello, J.P. (2014). UrbanSound: A dataset for urban sound classification. "
        "Proceedings of the 22nd ACM International Conference on Multimedia.",
        "Piczak, K.J. (2015). ESC: Dataset for Environmental Sound Classification. "
        "Proceedings of the 23rd ACM International Conference on Multimedia.",
        "He, K. et al. (2016). Deep Residual Learning for Image Recognition. CVPR.",
        "Sandler, M. et al. (2018). MobileNetV2: Inverted Residuals and Linear Bottlenecks. CVPR.",
        "McFee, B. et al. (2015). librosa: Audio and Music Analysis in Python. SciPy.",
        "Paszke, A. et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library. NeurIPS.",
        "Park, D.S. et al. (2019). SpecAugment: A Simple Data Augmentation Method for ASR. INTERSPEECH.",
        "Hugging Face (2024). DynamicSuperb/EnvironmentalSoundClassification_ESC50-Animals dataset.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # --- Appendix ---
    add_heading(doc, "Appendix A — Reproducibility", level=1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Python", env["python"]],
            ["PyTorch", env["torch"]],
            ["Platform", env["platform"]],
            ["CUDA available", str(env["cuda_available"])],
            ["GPU", env["cuda_device"]],
            ["Random seed", str(cfg["training"]["seed"])],
            ["Project root", str(PROJECT_ROOT)],
        ],
    )
    add_para(doc, "Hyperparameters (config/config.yaml):", bold=True)
    add_bullets(
        doc,
        [
            f"Audio: {cfg['audio']['sample_rate']} Hz, {cfg['audio']['duration_sec']} s, mono={cfg['audio']['mono']}",
            f"Spectrogram: n_fft={cfg['spectrogram']['n_fft']}, hop={cfg['spectrogram']['hop_length']}, n_mels={cfg['spectrogram']['n_mels']}",
            f"Training: batch={cfg['training']['batch_size']}, lr={cfg['training']['learning_rate']}, epochs={cfg['training']['epochs']}",
        ],
    )

    return doc


def build_development_notes() -> Document:
    doc = Document()
    doc.add_heading("Development Notes — Full Project", level=0)
    add_para(doc, "B9AI104 Deep Learning — Continuous Assessment 1 (CA1)")
    add_para(doc, "Student: Nadeesha Jayasuriya | Module: B9AI104 | Lecturer: Dr. Shahram Azizi Sazi")
    doc.add_paragraph()

    add_heading(doc, "Document Purpose", level=1)
    add_para(
        doc,
        "This separate Word document accompanies the main assignment report. It records project "
        "development steps, individual contribution, problems encountered, and estimated time "
        "of completion as required by the CA1 assessment brief.",
    )

    add_heading(doc, "1. Estimated Time of Completion", level=1)
    add_table(
        doc,
        ["Phase", "Task", "Hours", "Status"],
        [
            ["Setup", "Repository scaffold, config, folder structure", "4", "Complete"],
            ["Step 1", "Dataset acquisition & EDA (UrbanSound8K + ESC-50)", "6", "Complete"],
            ["Step 2", "Preprocessing pipeline & validation", "10", "Complete"],
            ["Step 3", "Model design & UrbanSound8K training (3 models)", "15", "Complete"],
            ["Step 4", "ESC-50 transfer learning & cross-domain analysis", "9", "Complete"],
            ["Step 5", "Streamlit deployment & verification", "6", "Complete"],
            ["Step 6", "Error analysis & inference benchmarking", "7", "Complete"],
            ["Step 8", "Final report writing", "8", "Complete"],
            ["Total estimated time", "", "65 hours", ""],
        ],
    )
    add_para(
        doc,
        "Note: Stage 1 UrbanSound8K training ran on CPU (~4.7 hours wall-clock for all three models). "
        "Steps 4–6 and inference benchmarking used an RTX 4060 Laptop GPU.",
    )

    add_heading(doc, "2. Project Development Timeline", level=1)
    add_bullets(
        doc,
        [
            "Week 1–2: Dataset download, EDA, preprocessing pipeline design",
            "Week 3–4: UrbanSound8K training (Custom CNN, ResNet50, MobileNetV2)",
            "Week 5: ESC-50 transfer learning and cross-domain evaluation",
            "Week 6: Error analysis, inference benchmarking, Streamlit deployment",
            "Week 7: CA1 ablations, final report, presentation, and demo preparation",
        ],
    )

    add_heading(doc, "3. Work Completed", level=1)
    add_bullets(
        doc,
        [
            "Built end-to-end Mel-spectrogram preprocessing for 8,932 urban + 200 ESC clips",
            "Trained Custom CNN, ResNet50, and MobileNetV2 on UrbanSound8K (fold-10 test)",
            "Ran ESC-50 transfer learning comparison (urban transfer, scratch, ImageNet-only)",
            "Generated confusion matrices, error case studies, and inference benchmarks",
            "Deployed Streamlit app with urban and animal classification modes",
            "Produced step reports (Steps 1–6) and final combined report",
            "Ran CA1 ablation studies (Model 3 head sizes, dropout comparison)",
        ],
    )
    add_heading(doc, "4. CA1 Compliance Notes", level=1)
    add_bullets(
        doc,
        [
            "Three models mapped to CA1 structure: Model 1 = conventional Custom CNN baseline; "
            "Model 2 = ResNet50 transfer; Model 3 = customised MobileNetV2 head + fine-tune",
            "Accuracy and macro recall reported for all models on fold-10 test",
            "Per-class precision/recall/F1 table included for MobileNetV2 in main report Section 5.1.1",
            "model.summary() screenshots included for all three models (main report Section 2.2–2.4)",
            "Hyperparameter ablations documented in main report Section 4.4–4.5",
            "PyTorch scripts equivalent to course 5-cell notebook workflow (notebooks/04_ca1_model_training.ipynb)",
        ],
    )
    add_heading(doc, "5. Problems Faced & Solutions", level=1)
    add_table(
        doc,
        ["Problem", "Solution"],
        [
            ["ESC-50 HF download (torchcodec missing)", "Used Audio(decode=False) + soundfile for loading"],
            ["PyTorch CPU-only install", "Installed torch 2.6.0+cu124 for RTX 4060 GPU"],
            ["Long CPU training (~4.7 hrs)", "Accepted valid CPU results; used GPU for Steps 4–6"],
            ["Urban transfer underperformed ImageNet on ESC-50", "Reported honestly; ImageNet selected for animal mode"],
            ["Streamlit not on PATH", "Run via python -m streamlit run app/streamlit_app.py"],
        ],
    )

    add_heading(doc, "6. Student Contribution (Individual Project)", level=1)
    add_para(
        doc,
        "This was completed as an individual project. I selected UrbanSound8K and the ESC-50 animal "
        "I selected UrbanSound8K and the ESC-50 animal subset as the CA1 datasets, designed the "
        "Mel-spectrogram-to-RGB preprocessing pipeline, and implemented the full PyTorch codebase "
        "(preprocessing, training, evaluation, and deployment). I trained and compared three CNN "
        "models on UrbanSound8K fold 10, extended the work with ESC-50 transfer-learning "
        "experiments, and ran CA1 hyperparameter ablations (Model 3 head sizes and Custom CNN "
        "dropout). I produced all reports, figures, error analysis, the Streamlit demo app, "
        "presentation slides, and prepared the live demo script for the 21 June presentation.",
    )

    add_heading(doc, "7. Reproducibility Commands", level=1)
    commands = [
        "python scripts/run_step1_eda.py",
        "python scripts/run_step2_preprocess.py",
        "python scripts/run_step3_train.py",
        "python scripts/run_ca1_ablation_studies.py",
        "python scripts/run_step4_esc50.py",
        "python scripts/run_step6_error_analysis.py",
        "python scripts/verify_step5_deployment.py",
        "python scripts/generate_model_summaries.py",
        "python scripts/generate_final_report.py",
        "python -m streamlit run app/streamlit_app.py",
    ]
    for cmd in commands:
        p = doc.add_paragraph(cmd, style="List Bullet")
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    return doc


def main() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    data = load_step_summaries()
    cfg = load_config()
    env = get_env_info()
    ablation = load_ablation_summary()

    cover = build_cover_sheet()
    cover_path = FINAL_DIR / "Cover_Sheet.docx"
    cover.save(cover_path)

    report = build_main_report(data, cfg, env, ablation)
    report_path = FINAL_DIR / "Final_Assignment_Report.docx"
    try:
        report.save(report_path)
    except PermissionError:
        alt_path = FINAL_DIR / "Final_Assignment_Report_updated.docx"
        report.save(alt_path)
        print(f"WARNING: {report_path.name} is open — saved to {alt_path}")
        report_path = alt_path

    dev_notes = build_development_notes()
    dev_path = FINAL_DIR / "Final_Development_Notes.docx"
    dev_notes.save(dev_path)

    manifest = {
        "title": "Deep Learning-Based Environmental Sound Classification",
        "generated": date.today().isoformat(),
        "outputs": {
            "cover_sheet": str(cover_path),
            "main_report": str(report_path),
            "development_notes": str(dev_path),
        },
        "environment": env,
        "best_urban_model": data["step3"]["best_model"],
        "best_urban_macro_f1": data["step3"]["best_macro_f1"],
        "best_esc50_run": data["step4"]["best_esc50_run"],
        "best_esc50_macro_f1": data["step4"]["best_esc50_macro_f1"],
    }
    with (FINAL_DIR / "final_report_manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)

    print("Final report package generated:")
    print(f"  Cover sheet:       {cover_path}")
    print(f"  Main report:       {report_path}")
    print(f"  Development notes: {dev_path}")
    print(f"  Manifest:          {FINAL_DIR / 'final_report_manifest.json'}")


if __name__ == "__main__":
    main()
