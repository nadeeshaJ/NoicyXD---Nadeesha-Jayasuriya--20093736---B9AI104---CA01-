"""Generate Step 1 report section and development notes for assignment.

Run: python scripts/generate_step1_report.py
Outputs: reports/step1/Step1_Report.docx, reports/step1/Step1_Development_Notes.docx
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = PROJECT_ROOT / "reports" / "figures" / "step1"
STEP1_DIR = PROJECT_ROOT / "reports" / "step1"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_figure(doc, path: Path, caption: str, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()


def build_step1_report(summary: dict, inventory_rows: list[dict]) -> Document:
    doc = Document()

    title = doc.add_heading("Step 1 — Dataset Acquisition & Exploratory Data Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Deep Learning (B9AI104) — Environmental Sound Classification via Mel-Spectrogram Images",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Introduction to Step 1", level=1)
    add_para(
        doc,
        "This section documents dataset acquisition, verification, exploratory data analysis (EDA), "
        "and class mapping for Stage 1 (UrbanSound8K) and Stage 2 (ESC-50 animal subset). "
        "This work supports the Dataset Implementation section of the main assignment report.",
    )

    add_heading(doc, "2. Dataset Inventory", level=1)
    table = doc.add_table(rows=1 + len(inventory_rows), cols=6)
    table.style = "Table Grid"
    headers = ["Dataset", "Clips", "Classes", "Duration (s)", "Sample Rate", "Split Strategy"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(inventory_rows):
        cells = table.rows[r_idx + 1].cells
        cells[0].text = row.get("dataset", "")
        cells[1].text = str(row.get("total_clips", ""))
        cells[2].text = str(row.get("num_classes", ""))
        cells[3].text = f"{row.get('duration_min_sec', '')}–{row.get('duration_max_sec', '')} (mean {row.get('duration_mean_sec', '')})"
        cells[4].text = str(row.get("sample_rate_hz", ""))
        cells[5].text = f"Test fold {row.get('test_fold', '')}" if "Urban" in row.get("dataset", "") else "80/20 stratified (planned)"
    doc.add_paragraph()

    add_heading(doc, "3. UrbanSound8K (Stage 1 — Main Dataset)", level=1)
    urban = summary.get("urbansound8k", {})
    add_para(doc, "Source: Salamon & Bello (2014). Urban environmental sounds in 10 classes.", bold=True)
    add_bullets(
        doc,
        [
            f"Total clips: {urban.get('total_clips', '8732')}",
            f"Classes: air conditioner, car horn, children playing, dog bark, drilling, engine idling, gun shot, jackhammer, siren, street music",
            f"Official evaluation: train on folds 1–9, test on fold {urban.get('test_fold', 10)}",
            f"Missing files on disk: {urban.get('missing_files', 0)}",
            "Purpose: main experiment — preprocessing, three-model training, evaluation, Streamlit urban mode",
        ],
    )

    add_heading(doc, "3.1 UrbanSound8K EDA Findings", level=2)
    add_bullets(
        doc,
        [
            "Class distribution is approximately balanced (~870 clips per class).",
            "Clip durations vary (typically under 4 seconds); preprocessing will pad/trim to 4.0 s.",
            "Fold distribution is balanced across classes, supporting official fold-based evaluation.",
            "Sample rates vary across source recordings; all audio will be resampled to 22,050 Hz.",
        ],
    )

    add_figure(
        doc,
        FIG_DIR / "urban_class_distribution.png",
        "Figure 1.1 — UrbanSound8K class distribution (clips per class).",
    )
    add_figure(
        doc,
        FIG_DIR / "urban_duration_distribution.png",
        "Figure 1.2 — UrbanSound8K clip duration distribution from metadata.",
    )
    add_figure(
        doc,
        FIG_DIR / "urban_fold_heatmap.png",
        "Figure 1.3 — UrbanSound8K class counts per fold (heatmap).",
    )
    add_figure(
        doc,
        FIG_DIR / "urban_sample_waveforms.png",
        "Figure 1.4 — Example waveforms per UrbanSound8K class.",
        width=Inches(6),
    )

    add_heading(doc, "4. ESC-50 Animal Subset (Stage 2 — Extension)", level=1)
    esc = summary.get("esc50_animals", {})
    add_para(
        doc,
        "Source: DynamicSuperb/EnvironmentalSoundClassification_ESC50-Animals (Hugging Face). "
        "Ten animal classes from the ESC-50 benchmark.",
        bold=True,
    )
    add_bullets(
        doc,
        [
            f"Total clips: {esc.get('total_clips', 200)} (20 per class)",
            "Classes: dog, rooster, pig, cow, frog, cat, hen, insects, sheep, crow",
            "Sample rate: 44,100 Hz; duration: 5.0 s per clip",
            "Purpose: cross-domain extension and transfer learning from UrbanSound8K-trained models",
            "Note: smaller than UrbanSound8K — transfer learning is required for Stage 2",
        ],
    )

    add_heading(doc, "4.1 ESC-50 Animals EDA Findings", level=2)
    add_bullets(
        doc,
        [
            "Perfectly balanced: 20 clips per class.",
            "Uniform 5-second duration and 44,100 Hz sample rate.",
            "Suitable for evaluating pipeline reusability across acoustic domains.",
            "Dog class overlaps semantically with UrbanSound8K dog_bark — useful for cross-domain discussion.",
        ],
    )

    add_figure(
        doc,
        FIG_DIR / "esc50_class_distribution.png",
        "Figure 1.5 — ESC-50 Animals class distribution.",
    )
    add_figure(
        doc,
        FIG_DIR / "esc50_duration_distribution.png",
        "Figure 1.6 — ESC-50 Animals clip duration distribution.",
    )
    add_figure(
        doc,
        FIG_DIR / "esc50_sample_waveforms.png",
        "Figure 1.7 — Example waveforms per ESC-50 animal class.",
        width=Inches(6),
    )

    add_heading(doc, "5. Class Mappings", level=1)
    add_para(
        doc,
        "Integer label mappings were saved for both datasets to ensure consistent training labels:",
    )
    add_bullets(
        doc,
        [
            "data/splits/urbansound8k/class_mapping.json",
            "data/splits/esc50_animals/class_mapping.json",
        ],
    )

    add_heading(doc, "6. Step 1 Conclusions", level=1)
    add_bullets(
        doc,
        [
            "Both datasets are verified, complete, and ready for preprocessing (Step 2).",
            "UrbanSound8K provides sufficient data for full three-model training and fold-10 evaluation.",
            "ESC-50 animal subset enables a credible cross-domain extension with transfer learning.",
            "EDA confirms preprocessing requirements: mono conversion, resampling, and fixed 4 s duration.",
        ],
    )

    add_heading(doc, "7. Bibliography (Step 1)", level=1)
    add_bullets(
        doc,
        [
            "Salamon, J. and Bello, J.P. (2014). UrbanSound: A dataset for urban sound classification. MIRUM Workshop.",
            "Piczak, K.J. (2015). ESC-50: Dataset for Environmental Sound Classification. ACM Multimedia.",
            "DynamicSuperb (2024). EnvironmentalSoundClassification_ESC50-Animals. Hugging Face Datasets.",
        ],
    )

    return doc


def build_development_notes_step1() -> Document:
    doc = Document()
    doc.add_heading("Development Notes — Step 1", level=0)
    add_para(doc, "Project: Environmental Sound Classification | Step 1: Dataset Acquisition & EDA")

    add_heading(doc, "Work Completed", level=1)
    add_bullets(
        doc,
        [
            "Downloaded and organized UrbanSound8K into data/raw/urbansound8k/",
            "Downloaded ESC-50 Animals from Hugging Face into data/raw/esc50/",
            "Cleaned raw data folders (removed readme, credits, .DS_Store, nested folders)",
            "Ran EDA scripts: class counts, durations, fold heatmap, waveforms",
            "Created class mapping JSON files for both datasets",
            "Generated dataset inventory CSV and Step 1 report with figures",
        ],
    )

    add_heading(doc, "Estimated Time", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Hours"
    rows = [
        ("Data download", "UrbanSound8K + HF ESC-50 Animals", "3"),
        ("Data cleanup", "Folder structure and config paths", "1"),
        ("EDA script", "Analysis code and figure generation", "4"),
        ("Report writing", "Step 1 Word document with figures", "2"),
        ("Total Step 1", "", "10"),
    ]
    for i, (a, b, c) in enumerate(rows, start=0):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c

    add_heading(doc, "Issues & Resolutions", level=1)
    add_bullets(
        doc,
        [
            "Nested UrbanSound8K folder — flattened to data/raw/urbansound8k/audio and metadata.",
            "HF audio decode required torchcodec — resolved by decoding with soundfile (decode=False).",
            "ESC-50 HF subset has 200 clips (not full 400) — documented; transfer learning planned for Stage 2.",
        ],
    )

    return doc


def main():
    summary_path = STEP1_DIR / "step1_summary.json"
    inventory_path = STEP1_DIR / "dataset_inventory.csv"

    with summary_path.open(encoding="utf-8") as f:
        summary = json.load(f)

    inventory_rows = []
    if inventory_path.exists():
        import pandas as pd

        inventory_rows = pd.read_csv(inventory_path).to_dict(orient="records")

    STEP1_DIR.mkdir(parents=True, exist_ok=True)

    report = build_step1_report(summary, inventory_rows)
    report.save(STEP1_DIR / "Step1_Dataset_EDA_Report.docx")

    notes = build_development_notes_step1()
    notes.save(STEP1_DIR / "Step1_Development_Notes.docx")

    print(f"Saved: {STEP1_DIR / 'Step1_Dataset_EDA_Report.docx'}")
    print(f"Saved: {STEP1_DIR / 'Step1_Development_Notes.docx'}")


if __name__ == "__main__":
    main()
