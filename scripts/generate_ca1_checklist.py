"""Generate CA1 pre-submission compliance checklist (Phase 10)."""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = PROJECT_ROOT / "reports" / "final"


def main() -> None:
    ablation_exists = (FINAL_DIR / "ca1_ablation_summary.json").exists()
    doc = Document()
    title = doc.add_heading("CA1 Pre-Submission Checklist", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Generated {date.today().strftime('%d %B %Y')} — cross-checked against CA01_Guidelines.docx")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    items = [
        ("Done", "3 models: Custom CNN (M1), ResNet50 (M2), MobileNetV2 (M3)"),
        ("Done" if ablation_exists else "Pending", "Hyperparameter experiments (head 64/128/256 + dropout)", "reports/final/ca1_ablation_summary.json"),
        ("You", "Dataset shown to / approved by lecturer"),
        ("Done", "Introduction includes 2–3 sentence CA1 brief (Final report Section 1.0)"),
        ("Done", "Architecture diagrams (Conv vs Flatten boundary)"),
        ("Done", "Images normalised; SpecAugment on train only"),
        ("Done", "Cross-entropy loss with integer labels (PyTorch CrossEntropyLoss)"),
        ("Done", "Training + validation curves saved; best epoch by val loss"),
        ("Done", "Parameter counts in report and benchmarks"),
        ("Done", "Accuracy AND macro recall reported (Section 5.1)"),
        ("Done", "Probabilities in Streamlit app (confidence bars)"),
        ("Done", "All class names listed (report Section 1.0.2 + slides)"),
        ("Done", "Early stopping explained (patience=5)"),
        ("Done", "Augmentation justified in report Section 4.7"),
        ("Done", "Main report + development notes + presentation"),
        ("You", "Cover sheet completed (name, ID, programme)"),
        ("You", "Student contribution section in dev notes (Section 5 — individual project)"),
        ("You", "Presentation prepared for 21 June 2026"),
        ("You", "Submit before 28 June 2026 + Moodle originality declaration"),
        ("You", "Local backup copy saved"),
    ]

    table = doc.add_table(rows=1 + len(items), cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Item"
    table.rows[0].cells[2].text = "Notes"
    for i, row in enumerate(items, start=1):
        if len(row) == 3:
            status, item, notes = row
        else:
            status, item = row
            notes = ""
        table.rows[i].cells[0].text = status
        table.rows[i].cells[1].text = item
        table.rows[i].cells[2].text = notes

    doc.add_paragraph()
    doc.add_heading("Key Files to Submit", level=1)
    for path in [
        "reports/final/Final_Assignment_Report.docx",
        "reports/final/Final_Development_Notes.docx",
        "reports/final/Cover_Sheet.docx",
        "reports/final/Presentation_Slides.pptx",
        "notebooks/04_ca1_model_training.ipynb",
        "Full code repository (noicy_XD/)",
    ]:
        p = doc.add_paragraph(path, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    out = FINAL_DIR / "CA1_PreSubmission_Checklist.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
