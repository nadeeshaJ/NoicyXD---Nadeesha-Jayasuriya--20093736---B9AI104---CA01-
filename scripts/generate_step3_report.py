"""Generate Step 3 model design and training report.

Run: python scripts/generate_step3_report.py
Outputs: reports/step3/Step3_Report.docx, reports/step3/Step3_Development_Notes.docx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = PROJECT_ROOT / "reports" / "figures" / "step3"
STEP3_DIR = PROJECT_ROOT / "reports" / "step3"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_figure(doc, path: Path, caption: str, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()


def build_report(summary: dict, comparison_rows: list[dict]) -> Document:
    doc = Document()
    title = doc.add_heading("Step 3 — Model Architecture, Design & UrbanSound8K Training", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Deep Learning (B9AI104) — Design & Architecture (25%) + Model Execution (25%)")
    doc.add_paragraph()

    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "Three deep learning models classify Mel-spectrogram RGB images of environmental sounds. "
        "UrbanSound8K fold-10 test evaluation compares a custom CNN baseline against "
        "ResNet50 and MobileNetV2 transfer learning models.",
    )

    add_heading(doc, "2. Model Architecture Background", level=1)

    add_heading(doc, "2.1 Custom CNN (Baseline)", level=2)
    add_bullets(
        doc,
        [
            "Role: demonstrate CNN design from scratch without pretrained weights",
            "Input: 224x224x3 Mel-spectrogram RGB image",
            "Conv blocks: 32 -> 64 -> 128 -> 128 filters with ReLU and MaxPooling",
            "Classifier: Flatten -> FC(256) -> Dropout(0.5) -> FC(10)",
            "Output: 10-class softmax (urban sound classes)",
            "Loss: Cross-entropy | Optimizer: Adam | LR: 1e-4",
        ],
    )
    add_figure(doc, FIG_DIR / "architecture_custom_cnn.png", "Figure 3.1 — Custom CNN architecture.")

    add_heading(doc, "2.2 ResNet50 (Transfer Learning)", level=2)
    add_bullets(
        doc,
        [
            "Role: strong pretrained feature extractor (ImageNet weights)",
            "Backbone: ResNet50 with residual skip connections",
            "Strategy: Phase 1 — freeze backbone, train FC head; Phase 2 — fine-tune top layers",
            "Input normalization: ImageNet mean/std",
            "Replaced final FC layer for 10 urban classes",
        ],
    )

    add_heading(doc, "2.3 MobileNetV2 (Efficient Transfer Learning)", level=2)
    add_bullets(
        doc,
        [
            "Role: lightweight model for faster training and deployment inference",
            "Uses depthwise separable convolutions — fewer parameters than ResNet50",
            "Same two-phase fine-tuning strategy as ResNet50",
            "Trade-off: speed and size vs maximum accuracy",
        ],
    )
    add_figure(
        doc,
        FIG_DIR / "architecture_transfer_learning.png",
        "Figure 3.2 — Transfer learning strategy for ResNet50 and MobileNetV2.",
    )

    add_heading(doc, "3. Model Rules Summary", level=1)
    rules_table = doc.add_table(rows=4, cols=4)
    rules_table.style = "Table Grid"
    headers = ["Rule", "Custom CNN", "ResNet50", "MobileNetV2"]
    for i, h in enumerate(headers):
        rules_table.rows[0].cells[i].text = h
    rules = [
        ("Input size", "224x224x3", "224x224x3", "224x224x3"),
        ("Pretrained weights", "No", "ImageNet", "ImageNet"),
        ("Regularization", "Dropout 0.5", "Dropout 0.3 (head)", "Dropout 0.2 (head)"),
    ]
    for r, row in enumerate(rules, start=1):
        for c, val in enumerate(row):
            rules_table.rows[r].cells[c].text = val
    doc.add_paragraph()

    add_heading(doc, "4. Training Protocol", level=1)
    add_bullets(
        doc,
        [
            "Dataset: UrbanSound8K Mel-spectrogram PNG images",
            "Train: folds 1-9 (7,105 clips) | Val: 10% stratified (790) | Test: fold 10 (837)",
            "Batch size: 32 | Optimizer: Adam | Learning rate: 1e-4",
            "Early stopping: patience 5 on validation loss",
            "Best checkpoint saved by lowest validation loss",
        ],
    )

    add_heading(doc, "5. UrbanSound8K Test Results", level=1)
    if comparison_rows:
        table = doc.add_table(rows=1 + len(comparison_rows), cols=5)
        table.style = "Table Grid"
        cols = ["Model", "Accuracy", "Macro F1", "Weighted F1", "Train Time (s)"]
        for i, h in enumerate(cols):
            table.rows[0].cells[i].text = h
        for r_idx, row in enumerate(comparison_rows):
            cells = table.rows[r_idx + 1].cells
            cells[0].text = row.get("model", "")
            cells[1].text = f"{row.get('accuracy', 0):.4f}"
            cells[2].text = f"{row.get('macro_f1', 0):.4f}"
            cells[3].text = f"{row.get('weighted_f1', 0):.4f}"
            cells[4].text = str(row.get("train_time_sec", ""))
        doc.add_paragraph()

    best_model = summary.get("best_model", "")
    add_para(doc, f"Best model by macro F1: {best_model}", bold=True)

    add_figure(
        doc,
        FIG_DIR / "model_comparison_urbansound8k.png",
        "Figure 3.3 — Model comparison on UrbanSound8K test set (fold 10).",
    )

    for model_name in ("custom_cnn", "resnet50", "mobilenetv2"):
        hist = FIG_DIR / "urbansound8k" / model_name / "training_history.png"
        cm = FIG_DIR / "urbansound8k" / model_name / "confusion_matrix_test.png"
        add_figure(doc, hist, f"Figure — {model_name} training history.")
        add_figure(doc, cm, f"Figure — {model_name} test confusion matrix.", width=Inches(5))

    add_heading(doc, "6. Design Discussion", level=1)
    add_bullets(
        doc,
        [
            "Custom CNN: simplest architecture, trains quickly, may underfit complex urban sounds",
            "ResNet50: residual connections enable deeper feature learning; typically highest accuracy",
            "MobileNetV2: best parameter efficiency; suitable for Streamlit deployment if accuracy is acceptable",
            "Transfer learning leverages ImageNet spatial features applicable to spectrogram textures",
        ],
    )

    add_heading(doc, "7. Step 3 Conclusions", level=1)
    add_bullets(
        doc,
        [
            "All three models trained successfully on UrbanSound8K with official fold-10 evaluation.",
            f"Best performing model: {best_model} — selected for ESC-50 transfer learning (Step 4).",
            "Training curves and confusion matrices saved for report and presentation.",
            "Models and checkpoints saved in experiments/urbansound8k/.",
        ],
    )

    add_heading(doc, "8. Bibliography", level=1)
    add_bullets(
        doc,
        [
            "He, K. et al. (2016). Deep Residual Learning for Image Recognition. CVPR.",
            "Sandler, M. et al. (2018). MobileNetV2: Inverted Residuals and Linear Bottlenecks. CVPR.",
            "Salamon, J. and Bello, J.P. (2014). UrbanSound: A dataset for urban sound classification.",
        ],
    )

    return doc


def build_dev_notes() -> Document:
    doc = Document()
    doc.add_heading("Development Notes — Step 3", level=0)
    add_para(doc, "Step 3: Model Design and UrbanSound8K Training")
    add_heading(doc, "Work Completed", level=1)
    add_bullets(
        doc,
        [
            "Implemented Custom CNN, ResNet50, MobileNetV2 in src/models/",
            "Built training loop with early stopping and two-phase fine-tuning",
            "Trained all three models on UrbanSound8K (fold-10 test evaluation)",
            "Generated confusion matrices, training curves, model comparison plots",
            "Saved best checkpoints to experiments/urbansound8k/",
            "Produced Step 3 report with architecture diagrams and rules tables",
        ],
    )
    add_heading(doc, "Estimated Time", level=1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Hours"
    for i, (a, b, c) in enumerate([
        ("Model code", "3 architectures + factory", "6"),
        ("Training pipeline", "train.py, evaluate.py", "4"),
        ("Training runs", "3 models UrbanSound8K", "4"),
        ("Figures", "Architecture + results plots", "2"),
        ("Report", "Step 3 Word document", "3"),
        ("Total Step 3", "", "19"),
    ]):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
    return doc


def main():
    with (STEP3_DIR / "step3_summary.json").open(encoding="utf-8") as f:
        summary = json.load(f)

    comparison_rows = []
    csv_path = STEP3_DIR / "model_comparison_urbansound8k.csv"
    if csv_path.exists():
        import pandas as pd
        comparison_rows = pd.read_csv(csv_path).to_dict(orient="records")

    report = build_report(summary, comparison_rows)
    report.save(STEP3_DIR / "Step3_Model_Training_Report.docx")
    build_dev_notes().save(STEP3_DIR / "Step3_Development_Notes.docx")
    print(f"Saved: {STEP3_DIR / 'Step3_Model_Training_Report.docx'}")


if __name__ == "__main__":
    main()
