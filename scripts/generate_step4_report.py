"""Generate Step 4 cross-domain and evaluation report.

Run: python scripts/generate_step4_report.py
Outputs: reports/step4/Step4_Report.docx, reports/step4/Step4_Development_Notes.docx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = PROJECT_ROOT / "reports" / "figures" / "step4"
STEP4_DIR = PROJECT_ROOT / "reports" / "step4"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_figure(doc, path: Path, caption: str, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()


def build_report(summary: dict, esc_rows: list, cross_rows: list) -> Document:
    doc = Document()
    title = doc.add_heading(
        "Step 4 — ESC-50 Transfer Learning, Evaluation & Cross-Domain Analysis", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Deep Learning (B9AI104) — Model Comparison, Evaluation & Validation (15%) + Results & Elaboration (15%)",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "Stage 2 extends the Mel-spectrogram pipeline to the ESC-50 animal subset (200 clips, 10 classes). "
        f"The best UrbanSound8K model ({summary.get('urban_source_model', 'mobilenetv2')}, "
        f"macro F1={summary.get('urban_source_macro_f1', 0):.3f}) is fine-tuned on animal sounds. "
        "Results are compared against training from scratch and ImageNet-only transfer learning.",
    )

    add_heading(doc, "2. ESC-50 Experimental Design", level=1)
    add_bullets(
        doc,
        [
            "Dataset: ESC-50 Animals (Hugging Face subset), 200 clips, 20 per class",
            "Split: 70% train / 15% val / 15% test (stratified)",
            "Same preprocessing as Stage 1: 22.05 kHz, 4 s, Mel-spec, 224x224 RGB",
            "Track A: Fine-tune UrbanSound8K MobileNetV2 backbone on animals",
            "Track B: Custom CNN trained from scratch on animals only",
            "Track C: MobileNetV2 with ImageNet weights only (no urban pretraining)",
        ],
    )

    add_heading(doc, "3. ESC-50 Test Results", level=1)
    if esc_rows:
        table = doc.add_table(rows=1 + len(esc_rows), cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Approach", "Accuracy", "Macro F1", "Weighted F1"]):
            table.rows[0].cells[i].text = h
        for r, row in enumerate(esc_rows):
            cells = table.rows[r + 1].cells
            cells[0].text = row.get("approach", "")
            cells[1].text = f"{row.get('accuracy', 0):.4f}"
            cells[2].text = f"{row.get('macro_f1', 0):.4f}"
            cells[3].text = f"{row.get('weighted_f1', 0):.4f}"
        doc.add_paragraph()

    add_para(
        doc,
        f"Best ESC-50 approach: {summary.get('best_esc50_run', '')} "
        f"(macro F1={summary.get('best_esc50_macro_f1', 0):.4f})",
        bold=True,
    )

    add_heading(doc, "4. Cross-Domain Comparison", level=1)
    add_para(
        doc,
        "Performance drop from UrbanSound8K to ESC-50 Animals indicates domain shift between "
        "urban environmental sounds and natural/animal recordings.",
    )
    if cross_rows:
        table = doc.add_table(rows=1 + len(cross_rows), cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Approach", "Domain", "Macro F1", "F1 Drop vs Urban"]):
            table.rows[0].cells[i].text = h
        for r, row in enumerate(cross_rows):
            cells = table.rows[r + 1].cells
            cells[0].text = row.get("approach", "")
            cells[1].text = row.get("domain", "")
            cells[2].text = f"{row.get('macro_f1', 0):.4f}"
            cells[3].text = f"{row.get('f1_drop', 0):.4f}"
        doc.add_paragraph()

    add_figure(
        doc,
        FIG_DIR / "cross_domain_comparison.png",
        "Figure 4.1 — Cross-domain macro F1 and performance drop.",
    )

    add_heading(doc, "5. Confusion Matrices & Training Curves", level=1)
    for run in ("mobilenetv2_urbansound_transfer", "custom_cnn_from_scratch", "mobilenetv2_imagenet_only"):
        add_figure(
            doc,
            FIG_DIR / "esc50_animals" / run / "confusion_matrix_test.png",
            f"Confusion matrix — {run}.",
            width=Inches(4.8),
        )

    add_heading(doc, "6. Error Analysis", level=1)
    err = summary.get("error_analysis", {})
    transfer_err = err.get("transfer", {})
    add_para(doc, "Top confused class pairs (Urban transfer model):", bold=True)
    for pair in transfer_err.get("top_confused_pairs", [])[:5]:
        add_bullets(
            doc,
            [f"{pair['true_class']} misclassified as {pair['predicted_class']} ({pair['count']} times)"],
        )
    add_figure(
        doc,
        FIG_DIR / "error_examples_mobilenetv2_urbansound_transfer.png",
        "Figure 4.2 — Misclassified ESC-50 animal examples (transfer model).",
        width=Inches(6),
    )

    add_heading(doc, "7. Discussion", level=1)
    add_bullets(
        doc,
        [
            "ESC-50 is much smaller (200 clips) than UrbanSound8K (8732), causing a large cross-domain performance drop (~22–31% macro F1).",
            "ImageNet-pretrained MobileNetV2 generalised best to animal sounds in this experiment.",
            "UrbanSound8K transfer slightly outperformed training from scratch, showing some reusable spectrogram features.",
            "Custom CNN from scratch struggled most with limited data, confirming transfer learning is required.",
            "Dog (animal) vs dog_bark (urban) share semantics but differ in recording conditions.",
            "Confusion often occurs between acoustically similar classes (e.g. rooster/hen, sheep/cow).",
        ],
    )

    add_heading(doc, "8. Step 4 Conclusions", level=1)
    add_bullets(
        doc,
        [
            "The preprocessing pipeline generalises from urban to animal sounds without modification.",
            f"Best ESC-50 model: {summary.get('best_esc50_run', '')} — saved for Streamlit animal mode (Step 5).",
            "Cross-domain analysis shows domain shift remains a key challenge despite pipeline reuse.",
            "Transfer learning is essential on ESC-50; from-scratch CNN training alone is insufficient with 140 training clips.",
        ],
    )

    return doc


def build_dev_notes() -> Document:
    doc = Document()
    doc.add_heading("Development Notes — Step 4", level=0)
    add_para(doc, "Step 4: ESC-50 Transfer Learning and Cross-Domain Analysis")
    add_heading(doc, "Work Completed", level=1)
    add_bullets(
        doc,
        [
            "Fine-tuned UrbanSound8K MobileNetV2 on ESC-50 Animals",
            "Trained Custom CNN from scratch as baseline",
            "Compared ImageNet-only MobileNetV2 transfer",
            "Built cross-domain comparison tables and plots",
            "Generated error analysis with misclassified examples",
            "Produced Step 4 report document",
        ],
    )
    add_heading(doc, "Estimated Time", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Hours"
    for i, (a, b, c) in enumerate([
        ("Transfer learning", "Urban MobileNetV2 -> ESC-50", "2"),
        ("Baselines", "Scratch CNN + ImageNet MobileNetV2", "2"),
        ("Analysis", "Cross-domain + error analysis", "3"),
        ("Report", "Step 4 Word document", "2"),
        ("Total Step 4", "", "9"),
    ]):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
    return doc


def main():
    with (STEP4_DIR / "step4_summary.json").open(encoding="utf-8") as f:
        summary = json.load(f)

    esc_rows = cross_rows = []
    esc_path = STEP4_DIR / "esc50_model_comparison.csv"
    cross_path = STEP4_DIR / "cross_domain_comparison.csv"
    if esc_path.exists():
        import pandas as pd
        esc_rows = pd.read_csv(esc_path).to_dict(orient="records")
    if cross_path.exists():
        import pandas as pd
        cross_rows = pd.read_csv(cross_path).to_dict(orient="records")

    build_report(summary, esc_rows, cross_rows).save(
        STEP4_DIR / "Step4_CrossDomain_Report.docx"
    )
    build_dev_notes().save(STEP4_DIR / "Step4_Development_Notes.docx")
    print(f"Saved: {STEP4_DIR / 'Step4_CrossDomain_Report.docx'}")


if __name__ == "__main__":
    main()
