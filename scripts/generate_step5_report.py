"""Generate Step 5 deployment report and development notes.

Run: python scripts/generate_step5_report.py
Outputs: reports/step5/Step5_Report.docx, reports/step5/Step5_Development_Notes.docx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
STEP5_DIR = PROJECT_ROOT / "reports" / "step5"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def build_report(summary: dict) -> Document:
    doc = Document()
    title = doc.add_heading("Step 5 — Application Deployment (Streamlit Web App)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Deep Learning (B9AI104) — Application Deployment Section")
    doc.add_paragraph()

    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "A Streamlit web application deploys the trained environmental sound classifiers for live inference. "
        "Users upload a WAV file and the app runs the same preprocessing pipeline used in training, "
        "visualises the waveform and Mel-spectrogram, and returns top-3 class predictions with confidence scores.",
    )

    add_heading(doc, "2. Application Architecture", level=1)
    add_bullets(
        doc,
        [
            "Frontend / UI: Streamlit (`app/streamlit_app.py`)",
            "Inference logic: `src/predict.py`",
            "Shared config: `config/config.yaml` (preprocessing + model checkpoints)",
            "Models: PyTorch MobileNetV2 checkpoints from Steps 3 and 4",
        ],
    )

    add_heading(doc, "3. Deployment Pipeline", level=1)
    add_numbered = [
        "User selects Urban or Animal mode",
        "User uploads `.wav` audio file",
        "Audio loaded, resampled to 22,050 Hz, padded/trimmed to 4 s",
        "Mel-spectrogram generated and converted to 224×224 RGB image",
        "Image normalised (ImageNet stats for MobileNetV2)",
        "Model inference on GPU if available, otherwise CPU",
        "Top-3 predictions displayed with confidence bars",
    ]
    for item in add_numbered:
        doc.add_paragraph(item, style="List Number")

    deploy = summary.get("deployment", {})
    add_heading(doc, "4. Deployed Models", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Mode"
    table.rows[0].cells[1].text = "Model"
    table.rows[0].cells[2].text = "Checkpoint"
    urban = deploy.get("urban", {})
    animal = deploy.get("animal", {})
    table.rows[1].cells[0].text = "Urban Sound Mode"
    table.rows[1].cells[1].text = urban.get("model_name", "")
    table.rows[1].cells[2].text = urban.get("checkpoint", "")
    table.rows[2].cells[0].text = "Animal Sound Mode"
    table.rows[2].cells[1].text = animal.get("model_name", "")
    table.rows[2].cells[2].text = animal.get("checkpoint", "")
    doc.add_paragraph()

    add_heading(doc, "5. UI Features", level=1)
    add_bullets(
        doc,
        [
            "Mode selector: Urban vs Animal sound classification",
            "WAV file upload",
            "Waveform plot of uploaded audio",
            "Mel-spectrogram visualisation",
            "224×224 RGB model input preview",
            "Top prediction with confidence percentage",
            "Top-3 predictions with progress bars",
            "Expandable full class probability list",
            "Sidebar showing preprocessing settings and active model",
        ],
    )

    add_heading(doc, "6. Verification Tests", level=1)
    for check in summary.get("checks", []):
        add_bullets(
            doc,
            [
                f"{check['mode']} mode — sample `{Path(check.get('sample_audio', '')).name}` "
                f"→ predicted `{check.get('prediction')}` ({check.get('confidence', 0):.1%}) "
                f"on {check.get('device', 'cpu')}",
            ],
        )

    add_heading(doc, "7. How to Run", level=1)
    add_para(doc, "From the project root:", bold=True)
    add_bullets(doc, ["pip install -r requirements.txt", "streamlit run app/streamlit_app.py"])

    add_heading(doc, "8. Step 5 Conclusions", level=1)
    add_bullets(
        doc,
        [
            "The full audio-to-prediction pipeline is deployed in an interactive web application.",
            "Urban and animal modes reuse the same preprocessing with domain-specific models.",
            "The app is suitable for live demonstration in the assignment presentation.",
            "GPU acceleration is used automatically when CUDA is available.",
        ],
    )

    return doc


def build_dev_notes() -> Document:
    doc = Document()
    doc.add_heading("Development Notes — Step 5", level=0)
    add_para(doc, "Step 5: Streamlit Application Deployment")
    add_heading(doc, "Work Completed", level=1)
    add_bullets(
        doc,
        [
            "Implemented src/predict.py inference pipeline",
            "Built Streamlit app with urban and animal modes",
            "Connected trained MobileNetV2 checkpoints from Steps 3 and 4",
            "Added waveform, Mel-spectrogram, and RGB visualisations",
            "Verified inference on sample test clips",
            "Produced Step 5 deployment report",
        ],
    )
    add_heading(doc, "Estimated Time", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Hours"
    for i, (a, b, c) in enumerate([
        ("Inference module", "predict.py", "3"),
        ("Streamlit UI", "app/streamlit_app.py", "4"),
        ("Testing", "verify_step5_deployment.py", "1"),
        ("Report", "Step 5 Word document", "2"),
        ("Total Step 5", "", "10"),
    ]):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
    return doc


def main():
    STEP5_DIR.mkdir(parents=True, exist_ok=True)
    with (STEP5_DIR / "step5_verification.json").open(encoding="utf-8") as f:
        summary = json.load(f)
    build_report(summary).save(STEP5_DIR / "Step5_Deployment_Report.docx")
    build_dev_notes().save(STEP5_DIR / "Step5_Development_Notes.docx")
    print(f"Saved: {STEP5_DIR / 'Step5_Deployment_Report.docx'}")


if __name__ == "__main__":
    main()
